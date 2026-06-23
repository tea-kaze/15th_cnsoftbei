import os
import re
import json
import numpy as np
from pypdf import PdfReader
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "uploads")
INDEX_FILE = os.path.join(os.path.dirname(os.path.dirname(__file__)), "knowledge_index.json")

os.makedirs(UPLOAD_DIR, exist_ok=True)

_documents: list[dict] = []
_vectorizer: TfidfVectorizer | None = None
_doc_vectors = None


def _load_index():
    global _documents, _vectorizer, _doc_vectors
    if os.path.exists(INDEX_FILE):
        with open(INDEX_FILE, "r", encoding="utf-8") as f:
            _documents = json.load(f)
        _rebuild_index()


def _save_index():
    with open(INDEX_FILE, "w", encoding="utf-8") as f:
        json.dump(_documents, f, ensure_ascii=False)


def _rebuild_index():
    global _vectorizer, _doc_vectors
    if not _documents:
        _vectorizer = None
        _doc_vectors = None
        return
    _vectorizer = TfidfVectorizer(analyzer="char", ngram_range=(1, 3), max_features=2000)
    _doc_vectors = _vectorizer.fit_transform([d["text"] for d in _documents])


def _tokenize(text: str) -> list[str]:
    return re.findall(r"[一-鿿\w]+", text.lower())


def parse_file(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".txt":
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    elif ext == ".pdf":
        reader = PdfReader(filepath)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    elif ext == ".docx":
        doc = Document(filepath)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    texts.append(" | ".join(cells))
        return "\n".join(texts)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def split_text(text: str, chunk_size: int = 500) -> list[str]:
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks = []
    current = ""
    for para in paragraphs:
        if len(current) + len(para) <= chunk_size:
            current = (current + "\n" + para).strip()
        else:
            if current:
                chunks.append(current)
            current = para
    if current:
        chunks.append(current)
    # Merge small chunks
    merged = []
    buf = ""
    for c in chunks:
        if len(buf) + len(c) <= chunk_size:
            buf = (buf + "\n" + c).strip()
        else:
            if buf:
                merged.append(buf)
            buf = c
    if buf:
        merged.append(buf)
    return merged if merged else [""]


def add_document(filepath: str, filename: str) -> int:
    text = parse_file(filepath)
    chunks = split_text(text)
    if not chunks:
        return 0
    global _documents
    _documents = [d for d in _documents if d["source"] != filename]
    for i, chunk in enumerate(chunks):
        _documents.append({"id": f"{filename}_{i}", "source": filename, "text": chunk})
    _rebuild_index()
    _save_index()
    return len(chunks)


def search(query: str, top_k: int = 5) -> list[str]:
    if not _documents or _vectorizer is None:
        return []
    query_vec = _vectorizer.transform([query])
    scores = cosine_similarity(query_vec, _doc_vectors)[0]
    top_indices = np.argsort(scores)[-top_k:][::-1]
    return [_documents[i]["text"] for i in top_indices if scores[i] > 0]


def list_documents() -> list[dict]:
    seen = {}
    for d in _documents:
        src = d["source"]
        if src not in seen:
            seen[src] = 0
        seen[src] += 1
    return [{"name": k, "chunks": v} for k, v in seen.items()]


def delete_document(filename: str) -> bool:
    global _documents
    before = len(_documents)
    _documents = [d for d in _documents if d["source"] != filename]
    if len(_documents) < before:
        _rebuild_index()
        _save_index()
        return True
    return False


# Load index on startup
_load_index()
